from __future__ import annotations

import csv
import json
import re
import shutil
import subprocess
from contextlib import nullcontext
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Literal

from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import xlrd
import xlwt

from app.documents.ocr import (
    PdfOcrProvider,
    assess_pdf_ocr_admission,
    create_configured_ocr_provider,
    inspect_pdf_for_ocr,
    normalize_ocr_provider_name,
)


SupportedFormat = Literal["pdf", "doc", "docx", "xls", "xlsx", "csv", "txt", "json"]


@dataclass
class ExtractedDocument:
    source_path: str
    source_format: str
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)
    sheet_names: list[str] = field(default_factory=list)
    metadata: dict[str, str] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    provider: str = ""
    confidence: float | None = None
    quality: str = ""
    trace: dict[str, Any] = field(default_factory=dict)


def detect_format(path: Path) -> str:
    return path.suffix.lower().lstrip(".")


def _normalize_cell(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value)
    return str(value).strip()


def _read_text_with_fallbacks(path: Path, *, encodings: tuple[str, ...]) -> str:
    last_error = ""
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except Exception as exc:
            last_error = f"{encoding}: {exc}"
    raise ValueError(f"Cannot read text file: {path}. Last error: {last_error}")


def _read_csv_rows(path: Path) -> list[list[str]]:
    raw = _read_text_with_fallbacks(path, encodings=("utf-8-sig", "utf-8", "cp1251", "latin-1"))
    delimiter_candidates = [",", ";", "\t"]
    best_rows: list[list[str]] = []
    best_len = -1
    for delimiter in delimiter_candidates:
        reader = csv.reader(raw.splitlines(), delimiter=delimiter)
        rows = [[_normalize_cell(cell) for cell in row] for row in reader]
        score = sum(len(row) for row in rows[:20])
        if score > best_len:
            best_rows = rows
            best_len = score
    if best_rows:
        return best_rows
    raise ValueError(f"Cannot detect CSV rows for: {path}")


def _read_xlsx_tables(path: Path) -> tuple[list[list[list[str]]], list[str]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    tables: list[list[list[str]]] = []
    sheet_names: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            rows.append([_normalize_cell(cell) for cell in row])
        tables.append(rows)
        sheet_names.append(sheet.title)
    workbook.close()
    return tables, sheet_names


def _read_xls_tables(path: Path) -> tuple[list[list[list[str]]], list[str]]:
    workbook = xlrd.open_workbook(path)
    tables: list[list[list[str]]] = []
    sheet_names: list[str] = []
    for sheet in workbook.sheets():
        rows: list[list[str]] = []
        for row_index in range(sheet.nrows):
            row_values = sheet.row_values(row_index)
            rows.append([_normalize_cell(cell) for cell in row_values])
        tables.append(rows)
        sheet_names.append(sheet.name)
    return tables, sheet_names


def _read_docx(path: Path) -> ExtractedDocument:
    doc = Document(path)
    text_lines: list[str] = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        block: list[list[str]] = []
        for row in table.rows:
            block.append([_normalize_cell(cell.text) for cell in row.cells])
        tables.append(block)
    text = "\n".join(text_lines).strip()
    return ExtractedDocument(
        source_path=str(path),
        source_format="docx",
        text=text,
        tables=tables,
        metadata={"paragraph_count": str(len(text_lines)), "table_count": str(len(tables))},
    )


def _extract_printable_strings(data: bytes) -> str:
    ascii_chunks = re.findall(rb"[A-Za-z0-9@._:/\\-]{4,}", data)
    cp1251_text = data.decode("cp1251", errors="ignore")
    cp1251_chunks = re.findall(r"[\u0400-\u04FFA-Za-z0-9@._:/\\-]{4,}", cp1251_text)
    utf16_text = data.decode("utf-16le", errors="ignore")
    utf16_chunks = re.findall(r"[\u0400-\u04FFA-Za-z0-9@._:/\\-]{4,}", utf16_text)
    lines: list[str] = []
    lines.extend(chunk.decode("ascii", errors="ignore") for chunk in ascii_chunks[:600])
    lines.extend(cp1251_chunks[:600])
    lines.extend(utf16_chunks[:600])
    dedup: list[str] = []
    seen: set[str] = set()
    for item in lines:
        value = item.strip()
        if not value or value in seen:
            continue
        seen.add(value)
        dedup.append(value)
    return "\n".join(dedup)


def _read_doc(path: Path) -> ExtractedDocument:
    warnings: list[str] = []
    text = ""

    antiword = shutil.which("antiword")
    if antiword:
        completed = subprocess.run(
            [antiword, str(path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        if completed.returncode == 0 and completed.stdout.strip():
            text = completed.stdout.strip()
        else:
            warnings.append(f"antiword failed: rc={completed.returncode}")
    else:
        warnings.append("antiword not found in PATH; using binary printable-string fallback for DOC.")

    if not text:
        data = path.read_bytes()
        text = _extract_printable_strings(data).strip()
        warnings.append("DOC parsed via binary printable-string fallback.")
        if not antiword:
            warnings.append("Install antiword for higher-quality DOC extraction.")

    return ExtractedDocument(
        source_path=str(path),
        source_format="doc",
        text=text,
        warnings=warnings,
        metadata={
            "antiword_available": str(bool(antiword)).lower(),
            "fallback_used": str(bool(not text or warnings)).lower(),
        },
    )


def _merge_trace(base_trace: dict[str, Any], extra_trace: dict[str, Any]) -> dict[str, Any]:
    result = dict(base_trace)
    for key, value in extra_trace.items():
        if key not in result:
            result[key] = value
            continue
        existing = result[key]
        if isinstance(existing, dict) and isinstance(value, dict):
            result[key] = _merge_trace(existing, value)
        else:
            result[key] = value
    return result


def _read_pdf(
    path: Path,
    *,
    enable_ocr: bool = True,
    ocr_provider: PdfOcrProvider | None = None,
    ocr_provider_name: str | None = None,
    ocr_trace_dir: Path | None = None,
    ocr_execution_context: Any | None = None,
) -> ExtractedDocument:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append((page.extract_text() or "").strip())
    text = "\n\n".join(chunk for chunk in pages if chunk).strip()
    scan_summary = inspect_pdf_for_ocr(path)
    requested_provider_name = normalize_ocr_provider_name(ocr_provider_name)
    ocr_admission = assess_pdf_ocr_admission(path, scan_summary=scan_summary)
    payload = ExtractedDocument(
        source_path=str(path),
        source_format="pdf",
        text=text,
        metadata={
            "pages": str(len(reader.pages)),
            "pages_with_text": str(scan_summary.pages_with_text),
            "image_pages": str(scan_summary.image_pages),
            "image_only_pages": str(scan_summary.image_only_pages),
            "ocr_applied": "false",
            "ocr_attempted": "false",
            "ocr_trigger_reason": scan_summary.trigger_reason,
        },
        trace={
            "pdf_scan": scan_summary.to_trace(),
            "ocr_admission": ocr_admission.to_trace(),
        },
    )
    if not enable_ocr or not scan_summary.should_run_ocr:
        return payload
    if not ocr_admission.should_attempt:
        payload.metadata["ocr_skip_reason"] = ocr_admission.reason
        payload.metadata["ocr_page_cap"] = str(ocr_admission.page_cap)
        payload.metadata["ocr_size_cap_bytes"] = str(ocr_admission.size_cap_bytes)
        payload.warnings.append(
            "OCR skipped for scan-only PDF outside speed budget "
            f"(reason={ocr_admission.reason}, pages={ocr_admission.page_count}, size_bytes={ocr_admission.size_bytes})."
        )
        payload.trace = _merge_trace(
            payload.trace,
            {
                "ocr": {
                    "triggered": True,
                    "provider_requested": requested_provider_name,
                    "provider_selected": "",
                    "status": "budget_skipped",
                    "skip_reason": ocr_admission.reason,
                    "admission": ocr_admission.to_trace(),
                }
            },
        )
        return payload

    provider = ocr_provider or create_configured_ocr_provider(requested_provider_name)
    if provider is None:
        if requested_provider_name:
            payload.warnings.append(
                f"OCR provider '{requested_provider_name}' is not configured; returning base PDF extraction result."
            )
        else:
            payload.warnings.append("PDF appears scan-only or without text layer, but no OCR provider is configured.")
        payload.trace = _merge_trace(
            payload.trace,
            {
                "ocr": {
                    "triggered": True,
                    "provider_requested": requested_provider_name,
                    "provider_selected": "",
                    "status": "not_configured",
                }
            },
        )
        return payload

    payload.metadata["ocr_attempted"] = "true"
    try:
        execution_context = ocr_execution_context() if callable(ocr_execution_context) else nullcontext()
        with execution_context:
            ocr_result = provider.extract_pdf(path, trace_dir=ocr_trace_dir)
    except Exception as exc:
        payload.warnings.append(f"OCR fallback failed via {provider.name}: {exc}")
        payload.trace = _merge_trace(
            payload.trace,
            {
                "ocr": {
                    "triggered": True,
                    "provider_requested": requested_provider_name or provider.name,
                    "provider_selected": provider.name,
                    "status": "failed",
                    "error": str(exc),
                }
            },
        )
        return payload

    payload.warnings.extend(ocr_result.warnings)
    payload.trace = _merge_trace(
        payload.trace,
        {
            "ocr": {
                "triggered": True,
                "provider_requested": requested_provider_name or provider.name,
                "provider_selected": ocr_result.provider,
                "status": "succeeded" if ocr_result.text.strip() else "empty",
                "provider_trace": ocr_result.trace,
            }
        },
    )
    if ocr_result.text.strip():
        payload.text = ocr_result.text.strip()
        payload.provider = ocr_result.provider
        payload.confidence = ocr_result.confidence
        payload.quality = ocr_result.quality
        payload.metadata["ocr_applied"] = "true"
        payload.metadata["ocr_provider"] = ocr_result.provider
        if ocr_result.confidence is not None:
            payload.metadata["ocr_confidence"] = str(ocr_result.confidence)
        if ocr_result.quality:
            payload.metadata["ocr_quality"] = ocr_result.quality
    else:
        payload.warnings.append(f"OCR fallback via {ocr_result.provider} returned empty text.")
    return payload


def _read_txt(path: Path) -> ExtractedDocument:
    text = _read_text_with_fallbacks(path, encodings=("utf-8-sig", "utf-8", "cp1251", "latin-1"))
    return ExtractedDocument(source_path=str(path), source_format="txt", text=text)


def _read_json(path: Path) -> ExtractedDocument:
    payload = json.loads(_read_text_with_fallbacks(path, encodings=("utf-8-sig", "utf-8")))
    if isinstance(payload, dict):
        text = str(payload.get("text", "")).strip()
        tables = payload.get("tables") if isinstance(payload.get("tables"), list) else []
        metadata = payload.get("metadata") if isinstance(payload.get("metadata"), dict) else {}
        warnings = payload.get("warnings") if isinstance(payload.get("warnings"), list) else []
        trace = payload.get("trace") if isinstance(payload.get("trace"), dict) else {}
        confidence_value = payload.get("confidence")
        confidence = float(confidence_value) if isinstance(confidence_value, (int, float)) else None
        return ExtractedDocument(
            source_path=str(path),
            source_format="json",
            text=text,
            tables=tables,
            metadata={str(k): str(v) for k, v in metadata.items()},
            warnings=[str(item) for item in warnings],
            provider=str(payload.get("provider") or ""),
            confidence=confidence,
            quality=str(payload.get("quality") or ""),
            trace=trace,
        )
    return ExtractedDocument(source_path=str(path), source_format="json", text=json.dumps(payload, ensure_ascii=False))


def extract_document(
    path: Path,
    *,
    enable_ocr: bool = True,
    ocr_provider: PdfOcrProvider | None = None,
    ocr_provider_name: str | None = None,
    ocr_trace_dir: Path | None = None,
    ocr_execution_context: Any | None = None,
) -> ExtractedDocument:
    fmt = detect_format(path)
    if fmt == "pdf":
        return _read_pdf(
            path,
            enable_ocr=enable_ocr,
            ocr_provider=ocr_provider,
            ocr_provider_name=ocr_provider_name,
            ocr_trace_dir=ocr_trace_dir,
            ocr_execution_context=ocr_execution_context,
        )
    if fmt == "docx":
        return _read_docx(path)
    if fmt == "doc":
        return _read_doc(path)
    if fmt == "xlsx":
        tables, sheet_names = _read_xlsx_tables(path)
        text = "\n".join(" | ".join(row) for table in tables for row in table if any(cell for cell in row))
        return ExtractedDocument(
            source_path=str(path),
            source_format="xlsx",
            text=text,
            tables=tables,
            sheet_names=sheet_names,
            metadata={"sheet_count": str(len(sheet_names))},
        )
    if fmt == "xls":
        tables, sheet_names = _read_xls_tables(path)
        text = "\n".join(" | ".join(row) for table in tables for row in table if any(cell for cell in row))
        return ExtractedDocument(
            source_path=str(path),
            source_format="xls",
            text=text,
            tables=tables,
            sheet_names=sheet_names,
            metadata={"sheet_count": str(len(sheet_names))},
        )
    if fmt == "csv":
        rows = _read_csv_rows(path)
        text = "\n".join(" | ".join(row) for row in rows if any(cell for cell in row))
        return ExtractedDocument(
            source_path=str(path),
            source_format="csv",
            text=text,
            tables=[rows],
            sheet_names=["csv"],
            metadata={"row_count": str(len(rows))},
        )
    if fmt == "txt":
        return _read_txt(path)
    if fmt == "json":
        return _read_json(path)
    raise ValueError(f"Unsupported input format: {fmt}")


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _doc_text(payload: ExtractedDocument) -> str:
    text = (payload.text or "").strip()
    if text:
        return text
    if payload.tables:
        lines = []
        for table in payload.tables:
            for row in table:
                if any(cell for cell in row):
                    lines.append(" | ".join(row))
        return "\n".join(lines)
    return ""


def _to_single_table(payload: ExtractedDocument) -> list[list[str]]:
    if payload.tables:
        return payload.tables[0]
    text = _doc_text(payload)
    if not text:
        return []
    return [[line] for line in text.splitlines()]


def write_txt(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    output_path.write_text(_doc_text(payload), encoding="utf-8")


def write_json(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    output_path.write_text(json.dumps(asdict(payload), ensure_ascii=False, indent=2), encoding="utf-8")


def write_csv(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    rows = _to_single_table(payload)
    with output_path.open("w", encoding="utf-8-sig", newline="") as fp:
        writer = csv.writer(fp, delimiter=";")
        writer.writerows(rows)


def write_xlsx(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    workbook = Workbook()
    if payload.tables:
        for index, table in enumerate(payload.tables):
            ws = workbook.active if index == 0 else workbook.create_sheet()
            title = payload.sheet_names[index] if index < len(payload.sheet_names) else f"Sheet{index + 1}"
            ws.title = title[:31]
            for row in table:
                ws.append(row)
    else:
        ws = workbook.active
        ws.title = "Data"
        for line in _doc_text(payload).splitlines():
            ws.append([line])
    workbook.save(output_path)
    workbook.close()


def write_xls(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    book = xlwt.Workbook()
    tables = payload.tables if payload.tables else [[[_doc_text(payload)]]]
    for index, table in enumerate(tables):
        title = payload.sheet_names[index] if index < len(payload.sheet_names) else f"Sheet{index + 1}"
        sheet = book.add_sheet(title[:31])
        for row_index, row in enumerate(table):
            for col_index, cell in enumerate(row):
                sheet.write(row_index, col_index, cell)
    book.save(str(output_path))


def write_docx(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    doc = Document()
    text = _doc_text(payload)
    if text:
        for line in text.splitlines():
            doc.add_paragraph(line)
    for table in payload.tables:
        if not table:
            continue
        rows_count = len(table)
        cols_count = max(len(row) for row in table)
        doc_table = doc.add_table(rows=rows_count, cols=cols_count)
        for row_index, row in enumerate(table):
            for col_index, value in enumerate(row):
                doc_table.cell(row_index, col_index).text = value
    doc.save(output_path)


def _escape_rtf(value: str) -> str:
    return value.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")


def write_doc(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    text = _doc_text(payload)
    rtf_lines = [r"{\rtf1\ansi\deff0{\fonttbl{\f0 Calibri;}}", r"\f0\fs22"]
    for line in text.splitlines():
        rtf_lines.append(_escape_rtf(line) + r"\par")
    rtf_lines.append("}")
    output_path.write_text("\n".join(rtf_lines), encoding="utf-8")


def write_pdf(payload: ExtractedDocument, output_path: Path) -> None:
    _ensure_parent(output_path)
    c = canvas.Canvas(str(output_path), pagesize=A4)
    _, height = A4
    text = c.beginText(40, height - 40)
    text.setLeading(14)
    text.setFont("Helvetica", 10)
    for line in _doc_text(payload).splitlines():
        safe = line[:160]
        text.textLine(safe)
        if text.getY() < 60:
            c.drawText(text)
            c.showPage()
            text = c.beginText(40, height - 40)
            text.setLeading(14)
            text.setFont("Helvetica", 10)
    c.drawText(text)
    c.save()


def convert_document(
    input_path: Path,
    output_path: Path,
    *,
    enable_ocr: bool = True,
    ocr_provider: PdfOcrProvider | None = None,
    ocr_provider_name: str | None = None,
    ocr_trace_dir: Path | None = None,
) -> ExtractedDocument:
    payload = extract_document(
        input_path,
        enable_ocr=enable_ocr,
        ocr_provider=ocr_provider,
        ocr_provider_name=ocr_provider_name,
        ocr_trace_dir=ocr_trace_dir,
    )
    output_fmt = detect_format(output_path)
    if output_fmt == "txt":
        write_txt(payload, output_path)
    elif output_fmt == "json":
        write_json(payload, output_path)
    elif output_fmt == "csv":
        write_csv(payload, output_path)
    elif output_fmt == "xlsx":
        write_xlsx(payload, output_path)
    elif output_fmt == "xls":
        write_xls(payload, output_path)
    elif output_fmt == "docx":
        write_docx(payload, output_path)
    elif output_fmt == "doc":
        write_doc(payload, output_path)
    elif output_fmt == "pdf":
        write_pdf(payload, output_path)
    else:
        raise ValueError(f"Unsupported output format: {output_fmt}")
    return payload


__all__ = [
    "ExtractedDocument",
    "SupportedFormat",
    "convert_document",
    "detect_format",
    "extract_document",
    "write_csv",
    "write_doc",
    "write_docx",
    "write_json",
    "write_pdf",
    "write_txt",
    "write_xls",
    "write_xlsx",
]
